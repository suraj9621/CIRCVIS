from docx import Document
from docx.shared import Pt

output_path = 'CIRCVIS_Project_Documentation.docx'

doc = Document()

doc.add_heading('CIRCVIS: Context-Aware Waste Classification for Circular Cities', level=0)

doc.add_paragraph('')

doc.add_paragraph('PRACTICAL REPORT', style='Title')

doc.add_paragraph('Name')
nd = doc.add_paragraph('Meher Jeevan')
doc.add_paragraph('Email')
doc.add_paragraph('meherjeevan55@gmail.com')
doc.add_paragraph('Module')
doc.add_paragraph('Computer Vision')
doc.add_paragraph('Question No.')
doc.add_paragraph('1')

doc.add_paragraph('Artificial Intelligence Programming Assistance (2025-2026)')
doc.add_paragraph('Guidance: Mr. Hemant Kumar, Mr. Saksham Mishra')
doc.add_paragraph('NSTI Kanpur')
doc.add_paragraph('15-05-2026')

doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    'CIRCVIS is a context-aware waste classification system designed to automate waste sorting with deep learning. '
    'The project combines image preprocessing, model inference, and a responsive web dashboard to classify seven waste classes '
    'for smart city and recycling applications.'
)

doc.add_heading('Aim', level=1)
doc.add_paragraph(
    'To build an end-to-end waste classification application using Python, OpenCV, NumPy, TensorFlow/Keras, FastAPI, and a frontend dashboard. '
    'The application should read images, preprocess them, run classification, and display results through a web interface.'
)

doc.add_heading('Problem Statement', level=1)
doc.add_paragraph(
    'Urban waste streams contain mixed materials, degraded objects, occluded parts, and irregular lighting. Existing classification systems often fail '
    'to provide reliable real-time sorting in these conditions. CIRCVIS aims to overcome these challenges by combining ensemble modeling, calibration, '
    'and sustainability impact reporting.'
)

doc.add_heading('Objectives', level=1)
doc.add_paragraph('1. Research and implement context-aware waste classification using deep learning.')
doc.add_paragraph('2. Create a data pipeline for dataset preparation and annotation merging.')
doc.add_paragraph('3. Train and deploy models for reliable waste prediction across seven classes.')
doc.add_paragraph('4. Build a FastAPI backend with prediction endpoints and dashboard analytics.')
doc.add_paragraph('5. Provide a user-friendly frontend for live demo, prediction, and impact visualization.')

doc.add_heading('Requirements', level=1)
doc.add_paragraph('Software Requirements')
doc.add_paragraph('Python 3.10+ or 3.12')
doc.add_paragraph('VS Code or any code editor')
doc.add_paragraph('FastAPI, Uvicorn, OpenCV, NumPy, Pillow, TensorFlow/Keras')

doc.add_paragraph('Python Libraries Required')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('fastapi')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('uvicorn')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('numpy')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('opencv-python')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('Pillow')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('tensorflow')
libraries = doc.add_paragraph(style='List Bullet')
libraries.add_run('python-docx (for documentation generation)')

doc.add_heading('System Overview', level=1)
doc.add_paragraph(
    'CIRCVIS is structured into three main components: data processing, backend model serving, and frontend visualization. '
    'The data module handles dataset organization and train/test split. The backend module exposes prediction and analytics APIs. '
    'The frontend module provides a user interface for uploads, live demo, and dashboard charts.'
)

doc.add_heading('Project Structure', level=1)
doc.add_paragraph(
    'The repository contains the following key directories and files:'
)
doc.add_paragraph('backend/app/main.py - FastAPI server and SPA routing')
doc.add_paragraph('backend/app/routers/predict.py - Prediction endpoints')
doc.add_paragraph('backend/app/routers/dashboard.py - Analytics and impact APIs')
doc.add_paragraph('backend/app/services/model_service.py - Model loading and inference service')
doc.add_paragraph('backend/app/utils/helpers.py - Image preprocessing and utility helpers')
doc.add_paragraph('frontend/ - Static website pages and JavaScript for demo/dashboard')
doc.add_paragraph('data/ - Scripts for ETL, training, evaluation, finetuning, and preprocessing')
doc.add_paragraph('models/ - Saved model weights and calibration metadata')

doc.add_heading('Implementation Steps', level=1)
doc.add_paragraph('Step 1: Data Preparation')
doc.add_paragraph(
    'Use data/etl.py to merge and organize datasets into seven waste classes, then generate train/val/test splits.'
)
doc.add_paragraph('Step 2: Model Training')
doc.add_paragraph(
    'Use data/train_models.py to train ResNet50, MobileNetV2, EfficientNetB0, and ensemble configurations. '
    'The resulting models are stored in the models/ directory.'
)
doc.add_paragraph('Step 3: Backend Setup')
doc.add_paragraph(
    'Start the FastAPI app via backend/app/main.py. The backend loads trained models from models/circvis_model.keras and serves prediction routes.'
)
doc.add_paragraph('Step 4: Frontend Deployment')
doc.add_paragraph(
    'Open the frontend pages using the FastAPI server: /, /demo.html, /dashboard.html, and /about.html. '
    'The front end consumes the backend APIs for predictions and analytics.'
)

doc.add_heading('API Endpoints', level=1)
doc.add_paragraph('POST /api/predict - Upload a single image for classification.')
doc.add_paragraph('POST /api/predict-base64 - Send a base64 image string for prediction.')
doc.add_paragraph('POST /api/predict-batch - Upload multiple images at once for batch inference.')
doc.add_paragraph('POST /api/predict-url - Provide an image URL to classify remotely.')
doc.add_paragraph('GET /api/models - Return available loaded models and status.')
doc.add_paragraph('GET /api/classes - Return the list of supported waste classes.')
doc.add_paragraph('GET /api/stats - Return dashboard metrics and class distribution.')
doc.add_paragraph('GET /api/impact - Return estimated sustainability impact of predictions.')
doc.add_paragraph('GET /health - Health check for service readiness.')

doc.add_heading('Key Features', level=1)
doc.add_paragraph('• Seven-class waste classification: Plastic, Organic, Metal, Paper/Cardboard, Glass, Textile, Miscellaneous.')
doc.add_paragraph('• Ensemble model support with fallback to a single optimized Keras model.')
doc.add_paragraph('• Real-time prediction API with upload, batch, base64, and URL support.')
doc.add_paragraph('• Dashboard analytics, sustainability metrics, and case-study endpoints.')

doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    'CIRCVIS demonstrates a complete AI system for waste classification, from dataset preprocessing to model serving and frontend visualization. '
    'The project is ready for further enhancement with additional dataset sources, model optimization, and deployment on edge devices.'
)

doc.save(output_path)
print(f'Generated {output_path}')
